from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# Adding picture
document.add_picture("profile.JPG", width=Inches(2.0))

# Asking information
name = input("What is your name? ")
speak("Salam " + name + " Chetori?")
phone_number = input("What is your phone number? ")
email = input("What is your email? ")

#Adding information
document.add_paragraph(name + " | " + phone_number + " | " + email)

# About me
document.add_heading("About Me")
about_me = input("Tell me about yourself? ")
document.add_paragraph(about_me)

#Work experience
document.add_heading("Work Experiences")
p = document.add_paragraph("")
company = input("Where do u work? ")
from_date = input("From date? ")
to_date = input("To date? ")
p.add_run(company + "\n").bold = True
p.add_run(from_date + " - " + to_date + "\n").italic = True
experience_detail = input("Express your work experience ")
p.add_run(experience_detail)

# Add more experiences
while True:
    has_more_experiences = input("Do you have more experience? Yes or No ")

    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph("")
        company = input("Where do u work? ")
        from_date = input("From date? ")
        to_date = input("To date? ")
        p.add_run(company + "\n").bold = True
        p.add_run(from_date + " - " + to_date + "\n").italic = True
        experience_detail = input("Express your work experience ")
        p.add_run(experience_detail)
    else:
        break

#Skills
document.add_heading("Skills")
p = document.add_paragraph("")
skill = input("Add a Skill ")
p.add_run(skill)
p.style = "List Bullet"

# Add more skills
while True:
    has_more_skills = input("Do you have more skills? Yes or No ")

    if has_more_skills.lower() == "yes":
        p = document.add_paragraph("")
        skill = input("Add a Skill ")
        p.add_run(skill)
        p.style = "List Bullet"
    else:
        break


document.save("cv.docx")



